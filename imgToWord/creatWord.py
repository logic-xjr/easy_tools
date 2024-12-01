import json
import requests
from docx import Document
from docx.shared import Cm, Inches, Mm, Pt
from io import BytesIO
from convert_to_jpeg import convert_to_jpeg

# 读取JSON文件
with open('data.json', 'r', encoding='utf-8') as file:
    data = json.load(file)

# 创建一个新的Word文档
doc = Document()

# 设置页边距
left_margin = Cm(1)  # 左边距1厘米
right_margin = Cm(1)  # 右边距1厘米
top_margin = Cm(1)  # 上边距1厘米
bottom_margin = Cm(1)  # 下边距1厘米

section = doc.sections[0]
section.left_margin = left_margin
section.right_margin = right_margin
section.top_margin = top_margin
section.bottom_margin = bottom_margin

# 页面总宽度（以英寸为单位）
total_page_width_in_inches = 8.5  # A4纸宽度大约为8.5英寸

# 计算除去页边距后的页面宽度
page_width_cm = total_page_width_in_inches * 2.54  # 将英寸转换为厘米
effective_page_width_cm = page_width_cm - (left_margin.inches + right_margin.inches) * 2.54
effective_page_width_emu = int(effective_page_width_cm * 360000)


# 遍历JSON数据
for item in data:
    # 添加code
    doc.add_paragraph(item['code'])
    print(item['code'])
    # 遍历图片URL并添加到Word文档
    for image_url in item['images']:
        response = requests.get(image_url)
        if response.status_code == 200:
            # 将图片保存到BytesIO对象中
            image_stream = BytesIO(response.content)
            # 添加图片，宽度设置为除去页边距后的页面宽度
            doc.add_picture(convert_to_jpeg(image_stream), width=effective_page_width_emu)
            doc.add_paragraph()  # 添加一个空段落作为图片和下一个code之间的分隔

# 保存Word文档
doc.save('output.docx')